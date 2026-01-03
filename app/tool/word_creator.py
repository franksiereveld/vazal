import os
import requests
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
from app.tool.base import BaseTool, ToolResult

class WordCreatorTool(BaseTool):
    name: str = "word_creator"
    description: str = """
    Creates a professional Word document (.docx) with formatted text, images, and tables.
    * Supports Headings (Level 1-3), Paragraphs, Bullet Points, and Numbered Lists.
    * Can insert images from URLs (automatically downloaded).
    * Can create tables with headers.
    * Saves the final document to the 'output/' folder.
    """
    parameters: dict = {
        "type": "object",
        "properties": {
            "filename": {
                "type": "string",
                "description": "The name of the output file (e.g., 'report.docx'). Will be saved in 'output/' folder.",
            },
            "title": {
                "type": "string",
                "description": "The main title of the document.",
            },
            "content": {
                "type": "array",
                "description": "List of content blocks to add to the document.",
                "items": {
                    "type": "object",
                    "properties": {
                        "type": {
                            "type": "string",
                            "enum": ["heading", "paragraph", "bullet", "number", "image", "table", "page_break"],
                            "description": "The type of content block.",
                        },
                        "text": {
                            "type": "string",
                            "description": "The text content (for headings, paragraphs, lists).",
                        },
                        "level": {
                            "type": "integer",
                            "description": "Heading level (1-3) or List indentation level (0-2). Default: 1 for headings, 0 for lists.",
                        },
                        "url": {
                            "type": "string",
                            "description": "URL of the image (for type='image').",
                        },
                        "caption": {
                            "type": "string",
                            "description": "Caption for the image (optional).",
                        },
                        "rows": {
                            "type": "array",
                            "items": {
                                "type": "array",
                                "items": {"type": "string"}
                            },
                            "description": "List of rows for the table (first row is header).",
                        }
                    },
                    "required": ["type"],
                },
            },
        },
        "required": ["filename", "content"],
    }

    def _download_image(self, url: str) -> str:
        """Downloads an image from a URL to the 'input/' folder."""
        try:
            input_dir = "input"
            if not os.path.exists(input_dir):
                os.makedirs(input_dir)

            filename = url.split("/")[-1].split("?")[0]
            if not filename or "." not in filename:
                import uuid
                filename = f"image_{uuid.uuid4().hex[:8]}.jpg"
            
            local_path = os.path.join(input_dir, filename)

            headers = {
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
            }
            response = requests.get(url, headers=headers, timeout=10)
            response.raise_for_status()
            
            # Verify image
            image_data = response.content
            img = Image.open(io.BytesIO(image_data))
            img.verify()
            
            with open(local_path, 'wb') as f:
                f.write(image_data)
            
            return local_path
        except Exception as e:
            print(f"⚠️ Failed to download image {url}: {e}")
            return None

    async def execute(self, filename: str, content: list, title: str = None) -> ToolResult:
        try:
            # Ensure output directory exists
            output_dir = "output"
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)
            
            output_path = os.path.join(output_dir, filename)
            doc = Document()

            # Add Title
            if title:
                doc.add_heading(title, 0)

            for block in content:
                block_type = block.get("type")
                text = block.get("text", "")
                
                if block_type == "heading":
                    level = block.get("level", 1)
                    doc.add_heading(text, level=level)
                
                elif block_type == "paragraph":
                    doc.add_paragraph(text)
                
                elif block_type == "bullet":
                    doc.add_paragraph(text, style='List Bullet')
                
                elif block_type == "number":
                    doc.add_paragraph(text, style='List Number')
                
                elif block_type == "page_break":
                    doc.add_page_break()
                
                elif block_type == "image":
                    url = block.get("url")
                    if url:
                        local_path = self._download_image(url)
                        if local_path:
                            doc.add_picture(local_path, width=Inches(6.0))
                            caption = block.get("caption")
                            if caption:
                                p = doc.add_paragraph(caption)
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                p.runs[0].font.italic = True
                                p.runs[0].font.size = Pt(9)
                                p.runs[0].font.color.rgb = RGBColor(100, 100, 100)

                elif block_type == "table":
                    rows = block.get("rows", [])
                    if rows:
                        num_cols = len(rows[0])
                        table = doc.add_table(rows=1, cols=num_cols)
                        table.style = 'Table Grid'
                        
                        # Header
                        hdr_cells = table.rows[0].cells
                        for i, header_text in enumerate(rows[0]):
                            hdr_cells[i].text = str(header_text)
                            # Make header bold
                            for paragraph in hdr_cells[i].paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
                        
                        # Data Rows
                        for row_data in rows[1:]:
                            row_cells = table.add_row().cells
                            for i, cell_text in enumerate(row_data):
                                if i < num_cols:
                                    row_cells[i].text = str(cell_text)

            doc.save(output_path)
            return ToolResult(output=f"Word document saved successfully to: {output_path}")

        except Exception as e:
            return ToolResult(error=f"Failed to create Word document: {e}")
